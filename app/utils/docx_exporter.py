"""
PM Digital Employee - DOCX Exporter
项目经理数字员工系统 - Word文档导出器

生成Word格式的报告文档。
"""

import io
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

from app.core.logging import get_logger

logger = get_logger(__name__)


class DocxExporter:
    """
    Word文档导出器.

    生成Word格式的报告文档。
    """

    def __init__(self) -> None:
        """初始化导出器."""
        self._document = None

    def export_weekly_report(
        self,
        report_data: Dict[str, Any],
    ) -> bytes:
        """
        导出周报为Word文档.

        Args:
            report_data: 周报数据

        Returns:
            bytes: 文档二进制内容
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.warning("python-docx not installed, returning placeholder")
            return self._create_placeholder_docx()

        doc = Document()

        # 标题
        title = doc.add_heading(
            f"项目周报 - {report_data.get('project_name', '')}",
            level=0,
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 报告周期
        period = doc.add_paragraph()
        period.add_run(
            f"报告周期：{report_data.get('week_start', '')} - {report_data.get('week_end', '')}"
        )
        period.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 生效日期
        date_para = doc.add_paragraph()
        date_para.add_run(f"生成日期：{datetime.now().strftime('%Y-%m-%d')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # 空行

        # 项目状态
        doc.add_heading("项目状态", level=1)
        status_para = doc.add_paragraph()
        status_para.add_run(f"项目状态：{report_data.get('project_status', '')}")
        status_para.add_run(f"\n整体进度：{report_data.get('progress', 0)}%")

        # 本周工作完成情况
        doc.add_heading("本周工作完成情况", level=1)
        completed_tasks = report_data.get("tasks_completed", [])
        if completed_tasks:
            for task in completed_tasks:
                doc.add_paragraph(task.get("name", ""), style="List Bullet")
        else:
            doc.add_paragraph("无")

        # 进行中任务
        doc.add_heading("进行中任务", level=1)
        in_progress_tasks = report_data.get("tasks_in_progress", [])
        if in_progress_tasks:
            for task in in_progress_tasks:
                para = doc.add_paragraph(style="List Bullet")
                para.add_run(f"{task.get('name', '')} - 进度: {task.get('progress', 0)}%")
        else:
            doc.add_paragraph("无")

        # 风险提示
        doc.add_heading("风险提示", level=1)
        risks = report_data.get("risks", [])
        if risks:
            for risk in risks:
                level = risk.get("level", "low")
                level_map = {"high": "🔴 高", "medium": "🟡 中", "low": "🟢 低"}
                para = doc.add_paragraph(style="List Bullet")
                para.add_run(f"[{level_map.get(level, level)}] {risk.get('description', '')}")
        else:
            doc.add_paragraph("无风险")

        # 下周计划
        doc.add_heading("下周工作计划", level=1)
        next_week_plan = report_data.get("next_week_plan", [])
        if next_week_plan:
            for plan in next_week_plan:
                doc.add_paragraph(plan, style="List Bullet")
        else:
            doc.add_paragraph("待规划")

        # 导出到字节
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    def export_wbs(
        self,
        wbs_data: Dict[str, Any],
    ) -> bytes:
        """
        导出WBS为Word文档.

        Args:
            wbs_data: WBS数据

        Returns:
            bytes: 文档二进制内容
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.warning("python-docx not installed, returning placeholder")
            return self._create_placeholder_docx()

        doc = Document()

        # 标题
        title = doc.add_heading(
            f"WBS工作分解结构 - {wbs_data.get('project_name', '')}",
            level=0,
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # 空行

        # WBS内容
        wbs_content = wbs_data.get("wbs_content", "")
        if wbs_content:
            # 解析并添加内容
            lines = wbs_content.split("\n")
            for line in lines:
                line = line.strip()
                if line.startswith("#"):
                    # 标题
                    level = line.count("#")
                    doc.add_heading(line.replace("#", "").strip(), level=min(level, 3))
                elif line.startswith("-") or line.startswith("*"):
                    # 列表项
                    doc.add_paragraph(line[1:].strip(), style="List Bullet")
                elif line:
                    # 普通段落
                    doc.add_paragraph(line)

        # 导出
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    def export_meeting_minutes(
        self,
        minutes_data: Dict[str, Any],
    ) -> bytes:
        """
        导出会议纪要为Word文档.

        Args:
            minutes_data: 会议纪要数据

        Returns:
            bytes: 文档二进制内容
        """
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.warning("python-docx not installed, returning placeholder")
            return self._create_placeholder_docx()

        doc = Document()

        # 标题
        title = doc.add_heading(
            f"会议纪要 - {minutes_data.get('title', '会议')}",
            level=0,
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 会议信息
        doc.add_paragraph(f"会议时间：{minutes_data.get('time', '')}")
        doc.add_paragraph(f"参会人员：{minutes_data.get('participants', '')}")
        doc.add_paragraph(f"记录日期：{datetime.now().strftime('%Y-%m-%d')}")

        doc.add_paragraph()  # 空行

        # 会议内容
        content = minutes_data.get("content", "")
        if content:
            lines = content.split("\n")
            for line in lines:
                line = line.strip()
                if line.startswith("#"):
                    level = line.count("#")
                    doc.add_heading(line.replace("#", "").strip(), level=min(level, 3))
                elif line.startswith("-") or line.startswith("*"):
                    doc.add_paragraph(line[1:].strip(), style="List Bullet")
                elif line:
                    doc.add_paragraph(line)

        # 导出
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    def _create_placeholder_docx(self) -> bytes:
        """创建占位文档（当python-docx未安装时）."""
        return b"Placeholder - python-docx not installed"


# 全局导出器实例
_docx_exporter: Optional[DocxExporter] = None


def get_docx_exporter() -> DocxExporter:
    """获取文档导出器."""
    global _docx_exporter
    if _docx_exporter is None:
        _docx_exporter = DocxExporter()
    return _docx_exporter


async def export_report(
    report_type: str,
    data: Dict[str, Any],
) -> bytes:
    """
    便捷函数：导出报告.

    Args:
        report_type: 报告类型
        data: 报告数据

    Returns:
        bytes: 文档内容
    """
    exporter = get_docx_exporter()

    if report_type == "weekly_report":
        return exporter.export_weekly_report(data)
    elif report_type == "wbs":
        return exporter.export_wbs(data)
    elif report_type == "meeting_minutes":
        return exporter.export_meeting_minutes(data)
    else:
        raise ValueError(f"Unknown report type: {report_type}")